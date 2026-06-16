"""
parser.py
=========
Text extraction utility module for PDF and DOCX files.
Uses PyMuPDF (fitz) for PDF files and python-docx for DOCX files.
"""

import io
import re

import docx  # python-docx
import fitz  # PyMuPDF


def clean_parsed_text(text: str) -> str:
    """
    Cleans raw extracted text by removing redundant newlines and spacing.

    Args:
        text (str): Raw text

    Returns:
        str: Cleaned text
    """
    # Replace multiple whitespace characters (including newlines) with a single space
    cleaned = re.sub(r"\s+", " ", text)
    return cleaned.strip()


def parse_pdf(file_bytes: bytes) -> str:
    """
    Extracts text from PDF file bytes.

    Args:
        file_bytes (bytes): Binary PDF data

    Returns:
        str: Extracted text
    """
    extracted_text = ""
    # Open PDF document from memory
    with fitz.open(stream=file_bytes, filetype="pdf") as doc:
        for page in doc:
            extracted_text += page.get_text()

    return extracted_text


def parse_docx(file_bytes: bytes) -> str:
    """
    Extracts text from DOCX file bytes.

    Args:
        file_bytes (bytes): Binary DOCX data

    Returns:
        str: Extracted text
    """
    # Open Word document from memory
    doc = docx.Document(io.BytesIO(file_bytes))
    paragraphs_text = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            paragraphs_text.append(paragraph.text)

    # Include table text if available in the resume
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs_text.append(cell.text)

    return "\n".join(paragraphs_text)


def extract_text(file_bytes: bytes, content_type: str) -> str:
    """
    Extracts and cleans text from PDF or DOCX file bytes based on content-type.
    """
    if "pdf" in content_type.lower():
        raw_text = parse_pdf(file_bytes)
    else:
        raw_text = parse_docx(file_bytes)
    return clean_parsed_text(raw_text)
