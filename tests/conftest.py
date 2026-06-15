import os
import sys
from io import BytesIO
from unittest.mock import MagicMock, patch

import pytest

# ── Path setup ────────────────────────────────────────────────────────────────
ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, ROOT)
sys.path.insert(0, os.path.join(ROOT, "backend"))  # so `import parser` works

# ── Shared test data ──────────────────────────────────────────────────────────
ALL_ROLES = [
    "Python Developer", "Backend Developer", "Java Developer",
    "Data Scientist",   "Data Analyst",      "Data Engineer",
    "DevOps Engineer",  "Cloud Architect",
    "Frontend Developer", "Mobile Developer", "Web Developer",
    "Security Analyst", "Business Analyst",  "QA Engineer",
]

SAMPLE_RESUME = (
    "John Doe — Senior Python Developer\n"
    "5+ years building production-grade backend APIs with Django REST Framework, "
    "Flask, and FastAPI. Expert in PostgreSQL, Redis, Docker, Kubernetes, and AWS. "
    "Implemented CI/CD pipelines with GitHub Actions, designed microservices handling "
    "1M+ daily requests, and led squads of 5 engineers using agile methodology. "
    "Experienced in test-driven development, REST API design, and cloud deployments."
)

MOCK_ALL_PROBS = {role: 0.01 for role in ALL_ROLES}
MOCK_ALL_PROBS["Python Developer"] = 0.85

MOCK_PREDICT_RETURN = {
    "label": "Python Developer",
    "confidence": 0.85,
    "top3": [
        {"label": "Python Developer",  "score": 0.85},
        {"label": "Backend Developer", "score": 0.06},
        {"label": "Data Engineer",     "score": 0.03},
    ],
    "all_probs": MOCK_ALL_PROBS,
}

MOCK_DNA_RETURN = {
    "cluster_scores": {
        "Backend & Python":  30.33,
        "Data Science":       0.33,
        "DevOps & Cloud":     0.50,
        "Frontend & Mobile":  0.33,
        "Security":           1.00,
        "Business":           0.50,
    },
    "skill_gap": {
        "role": "Python Developer",
        "present": ["python", "django", "flask", "docker", "aws"],
        "missing": ["kubernetes", "terraform", "ansible", "jenkins",
                    "prometheus", "grafana", "linux", "bash", "go", "rust"],
        "fit_pct": 33.33,
    },
    "alternative_paths":  [
        {"role": "Backend Developer",
         "score": 0.06,
         "gap_count": 3},
    ],
}


# ── Session-wide ML mock ───────────────────────────────────────────────────────
# autouse=True means this runs automatically before EVERY test in the session.
# It prevents any test from accidentally loading real 80MB model files.
@pytest.fixture(scope="session", autouse=True)
def mock_ml_models():
    with patch("src.predict_v2._load_assets"), \
         patch("src.predict_v2._encoder",       new=MagicMock()), \
         patch("src.predict_v2._classifier",    new=MagicMock()), \
         patch("src.predict_v2._label_encoder", new=MagicMock()):
        yield


# ── FastAPI test client ────────────────────────────────────────────────────────
@pytest.fixture(scope="module")
def client(mock_ml_models):
    from fastapi.testclient import TestClient
    import backend.main as main

    # Disable rate limiting for testing
    main.app.state.limiter.enabled = False

    with patch.object(main, "_load_assets"),  \
         patch.object(main, "predict",          return_value=MOCK_PREDICT_RETURN), \
         patch.object(main, "is_models_loaded", return_value=True), \
         patch.object(main, "get_resume_dna",   return_value=MOCK_DNA_RETURN):
        with TestClient(main.app) as c:
            yield c


# ── Shared text fixture ───────────────────────────────────────────────────────
@pytest.fixture(scope="session")
def sample_text():
    return SAMPLE_RESUME


# ── PDF fixture (created in memory with PyMuPDF) ─────────────────────────────
@pytest.fixture(scope="session")
def pdf_bytes():
    import fitz
    doc  = fitz.open()
    page = doc.new_page()
    page.insert_text((50, 50), SAMPLE_RESUME, fontsize=11)
    data = doc.tobytes()
    doc.close()
    return data


# ── DOCX fixture (created in memory with python-docx) ────────────────────────
@pytest.fixture(scope="session")
def docx_bytes():
    from docx import Document
    doc = Document()
    doc.add_heading("John Doe — Senior Python Developer", level=1)
    doc.add_paragraph(SAMPLE_RESUME)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
